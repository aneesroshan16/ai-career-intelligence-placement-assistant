"""
Raw text extraction from uploaded resume files. Kept separate from
ner_pipeline.py (structured extraction) so the extraction backend
(pdfplumber/python-docx today) can change independently of the NER logic.
"""
from __future__ import annotations

import io

from app.core.exceptions import ResumeParseError


def extract_text(file_bytes: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return _extract_pdf(file_bytes)
    if file_type == "docx":
        return _extract_docx(file_bytes)
    raise ResumeParseError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        text = "\n".join(text_parts).strip()
        if not text:
            raise ResumeParseError("No extractable text found in PDF (it may be a scanned image).")
        return text
    except ResumeParseError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ResumeParseError(f"Failed to parse PDF: {exc}") from exc


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx

        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs).strip()
        if not text:
            raise ResumeParseError("No extractable text found in DOCX.")
        return text
    except ResumeParseError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ResumeParseError(f"Failed to parse DOCX: {exc}") from exc
