from docx import Document
from docx.shared import Inches, Pt


document = Document()
section = document.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

normal = document.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10)

title = document.add_paragraph()
run = title.add_run("Test Candidate")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(18)

document.add_paragraph("test.candidate@example.invalid | Software Engineering Student")

for heading, lines in {
    "Summary": ["Synthetic resume used only to verify the production document-upload pipeline."],
    "Skills": ["Python, TypeScript, React, FastAPI, SQL"],
    "Education": ["B.Tech, Computer Science - Expected 2027"],
    "Projects": ["Career Intelligence Dashboard - Built a role-readiness application using React and Python."],
}.items():
    paragraph = document.add_paragraph()
    heading_run = paragraph.add_run(heading)
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    for line in lines:
        document.add_paragraph(line, style="List Bullet" if heading == "Skills" else None)

document.core_properties.title = "Synthetic Production Upload Test Resume"
document.core_properties.author = "Production Verification"
document.save("production-upload-test-resume.docx")
